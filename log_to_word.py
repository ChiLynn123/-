import json
import docx
from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.shared import Pt,RGBColor
doc = docx.Document()


#设置字体格式
doc.styles['Normal'].font.name = u'宋体'
doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
doc.styles['Normal'].font.size = Pt(10.5)
doc.styles['Normal'].font.color.rgb = RGBColor(0,0,0)


class statis():
    def log_satis(self, start_time, end_time):  # 起始时间:结果时间,结束时间:结果时间
        userlist = []
        tasklist = []
        tmp_user = {}
        statis_result = {
            "users": {},
            "totaljudgeDict": {},
            "totalurlNum": 0,
            "totaluserNum": 0,
            "totaltaskNum": 0,
            "totalcrawlNum": 0,
            "totalJudgeNum": 0,
            "totalIpNum": 0
        }
        wfd_user = ["X7007", "x7007", "x7006", "X7006", "X6982", "x6982","X4122", "x4122", "x2813", "X2813", "X3007"]
        with open('res.log', 'r', encoding='utf-8') as f:
            for line in f:
                taskdict = json.loads(line)
                userid = taskdict["userId"]
                if userid[0] == 'T':
                    userid = "Test"
                if userid[0] != 'T' and userid not in wfd_user and int(taskdict["task_time"]) > int(start_time) and int(
                        taskdict["result_time"]) < int(end_time):
                    taskid = taskdict["taskID"]
                    urlNum = taskdict["urlNum"]
                    crawlNum = taskdict["crawlNum"]
                    judgeNum = taskdict["judgeNum"]
                    ipNum = taskdict["ipNum"]
                    judge_dict = taskdict["judgeDict"]
                    if taskid not in tasklist:
                        tasklist.append(taskid)
                        # 统计用户dict
                        if tmp_user.get(userid) is not None:
                            tmp_user[userid] += 1
                        else:
                            tmp_user[userid] = 1
                        if userid not in userlist:
                            userlist.append(userid)
                            user_dict = {
                                "urlNum": 0,
                                "crawlNum": 0,
                                "judgeNum": 0,
                                "ipNum": 0,
                                "judgeDict": {},  # '0020007': 5
                                "taskNum": 0,
                            }
                            statis_result["users"][userid] = user_dict

                            statis_result["users"][userid]["urlNum"] = urlNum
                            statis_result["users"][userid]["crawlNum"] = crawlNum
                            statis_result["users"][userid]["judgeNum"] = judgeNum
                            statis_result["users"][userid]["ipNum"] = ipNum
                            statis_result["users"][userid]["judgeDict"] = judge_dict
                            statis_result["users"][userid]["taskNum"] = 1

                        else:
                            if crawlNum == None:
                                crawlNum = 0
                            if judgeNum == None:
                                judgeNum = 0
                            if ipNum == None:
                                ipNum = 0
                            statis_result["users"][userid]["urlNum"] += urlNum
                            statis_result["users"][userid]["crawlNum"] += crawlNum
                            statis_result["users"][userid]["judgeNum"] += judgeNum
                            statis_result["users"][userid]["ipNum"] += ipNum
                            statis_result["users"][userid]["taskNum"] += 1
                            # 遍历,按key累加 #统计判别类别dict
                            if judge_dict is not None:
                                for key in judge_dict:
                                    if statis_result["users"][userid]["judgeDict"].get(key) is not None:
                                        statis_result["users"][userid]["judgeDict"][key] += judge_dict[key]
                                    else:
                                        statis_result["users"][userid]["judgeDict"][key] = judge_dict[key]
                            # judge_dict累加
                        # total累加
                        statis_result["totalurlNum"] += urlNum
                        statis_result["totalcrawlNum"] += crawlNum
                        statis_result["totalJudgeNum"] += judgeNum
                        statis_result["totalIpNum"] += ipNum

                        # ("ju:{}".format(judge_dict))
                        if judge_dict is not None:
                            for key in judge_dict:
                                if statis_result["totaljudgeDict"].get(key) is not None:
                                    statis_result["totaljudgeDict"][key] += judge_dict[key]
                                else:
                                    statis_result["totaljudgeDict"][key] = judge_dict[key]
        statis_result["totaluserNum"] = len(userlist)
        statis_result["totaltaskNum"] = len(tasklist)

        statis_result = json.dumps(statis_result)
        return statis_result

# 类别
def pipei(startdata, enddata):
    a = statis()
    res_dict_1 = a.log_satis(startdata, enddata)
    res_dict = json.loads(res_dict_1)
    # 类别编号字典
    with open("classdic.txt", "r", encoding="utf-8") as f:
        for line in f:
            clssdict = json.loads(line)
    # 工号字典
    with open("userdic.txt", "r", encoding="utf-8") as f:
        for line in f:
            codedict = json.loads(line)
    # res_dict["users"]映射
    userdict = {}
    for userid in res_dict["users"]:
        if userid in codedict:
            user = codedict[userid]
        else:
            user = userid
        userdict[user] = {
            # "urlNum": 0,
            # "crawlNum": 0,
            # "judgeNum": 0,
            # "ipNum": 0,
            # "taskNum": 0,
            "下发次数": 0,
            "下发总量": 0,
            "有ip的总量": 0,
            "爬虫总量": 0,
            "识别网站数量": 0,
            # "judgeDict": {},
            "识别网站类别": {},
        }
        userdict[user]["下发次数"] = res_dict["users"][userid]["taskNum"]
        userdict[user]["下发总量"] = res_dict["users"][userid]["urlNum"]
        userdict[user]["有ip的总量"] = res_dict["users"][userid]["ipNum"]
        userdict[user]["爬虫总量"] = res_dict["users"][userid]["crawlNum"]
        userdict[user]["识别网站数量"] = res_dict["users"][userid]["judgeNum"]
        judgedict = res_dict["users"][userid]["judgeDict"]
        for judgeid in judgedict:
            judgename = clssdict[judgeid]
            userdict[user]["识别网站类别"][judgename] = judgedict[judgeid]
    # res_dict["totaljudgeDict"]映射
    totaljudge_dict = {}
    for judgeid in res_dict["totaljudgeDict"]:
        judge_name = clssdict[judgeid]
        totaljudge_dict[judge_name] = res_dict["totaljudgeDict"][judgeid]
    print(res_dict)
    print(totaljudge_dict)
    print(userdict)
    return res_dict, totaljudge_dict, userdict


def dict_sum_pro(dict1,dict2):
    """
    输入本周的使用情况
    输出历史汇总的dict_sum和sum拼接str，并更新dict_sum保存
    """
    #读取数据
    with open("./dict_sum.txt", "r", encoding="utf-8") as f:
        for line in f:
            dict_sum = json.loads(line)
    #处理数据 合并为新的dict_sum
    for key,value in dict2.items():
        if key in dict_sum:
            dict_sum[key] += value
        else:
            dict_sum[key] = value
    dict_sum['totaltaskNum'] = dict_sum['totaltaskNum']+dict1['totaltaskNum']
    dict_sum['totalurlNum'] = dict_sum['totalurlNum']+dict1['totalurlNum']
    dict_sum['totalcrawlNum'] = dict_sum['totalcrawlNum']+dict1['totalcrawlNum']
    dict_sum['totalJudgeNum'] = dict_sum['totalJudgeNum']+dict1['totalJudgeNum']
    #
    res_sum=[]
    for i in dict_sum.keys():
        pj = f"{i}{dict_sum[i]}"
        res_sum.append(pj)
    result_sum = '、'.join(res_sum)
    #保存更新
    with open('dict_sum.txt', 'w', encoding='utf-8') as file:
        file.write(json.dumps(dict_sum, ensure_ascii=False))
    return result_sum, dict_sum


#生成word 输入dict1,2,3
def log_to_word(startdata, enddata):
    time_ = f'{startdata}~{enddata}'
    dict1, dict2, dict3=pipei(startdata, enddata)
    result_sum, dict_sum = dict_sum_pro(dict1, dict2)
    # 信息处理
    res = []
    for i in dict2.keys():
        pj = f"{i}{dict2[i]}"
        res.append(pj)
    result = '、'.join(res)  # 本周汇总
    para1 = '    {}时间段，网站侦察兵云服务版总计使用人数{}人，使用频次{}次、涉及{}个地市研发中心，下发总量{}，其中有ip的总量为{}、爬虫总量{}、识别总量{}，其中{}。'.format(time_,
        dict1['totaluserNum'], dict1['totaltaskNum'], 'xx', dict1['totalurlNum'], dict1['totalIpNum'], dict1['totalcrawlNum'],dict1['totalJudgeNum'],result)
    para2 = '    {}时间段，使用人数{}人，使用人次{}次，包括非IAO部门使用{}人、使用{}次，各地使用详情：'.format(time_, dict1['totaluserNum'],dict1['totaltaskNum'], 'xx', 'xx')
    para3 = '    自网站侦察兵云服务上线以来(2020.12.31)，累计使用人次{}次，涉及35个地市，下发URL总量{}、爬虫总量{}、识别总量{}，其中V1.1-V2.0版本累计识别{}。'.format(
        dict_sum['totaltaskNum'], dict_sum["totalurlNum"], dict_sum["totalcrawlNum"], dict_sum["totalJudgeNum"],result_sum)
    # 生成word周报
    # 标题
    Head0 = doc.add_heading("", level=0)  # 这里不填标题内容
    run = Head0.add_run("网站侦察兵运营周报")
    run.font.name = u'黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')
    run.font.color.rgb = RGBColor(0, 0, 0)
    #doc.add_heading('网站侦察兵运营周报', level=0)
    # 标题1
    Head1 = doc.add_heading('', level=1)
    run = Head1.add_run("1、概述")
    run.font.name = u'黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')
    run.font.color.rgb = RGBColor(0, 0, 0)

    doc.add_paragraph(para1)
    # paraObjl.add_run('This text is being added to the second paragraph.')
    # paragraph_format = paraObjl.paragraph_format
    # paragraph_format.left_indent = Inches(0.5)
    # 标题2
    run = doc.add_heading('', level=1).add_run(u"2、阶段使用详情")  # 应用场景示例标题
    run.font.name = u'黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')
    run.font.color.rgb = RGBColor(0, 0, 0)
    #doc.add_heading('2、阶段使用详情', level=1)
    doc.add_paragraph(para2)
    count = len(dict3)
    # 生成count+1行、2列的表格
    table = doc.add_table(rows=count + 1, cols=2)
    row = table.rows[0]
    row.cells[0].text = '用户'  # '第1行第一列'
    row.cells[1].text = '下发详情'  # '第1行第二列'
    cou = 0
    for i in dict3.keys():
        cou += 1
        row = table.rows[cou]
        row.cells[0].text = i
        row.cells[1].text = str(dict3[i])

    # 标题3
    run = doc.add_heading('', level=1).add_run(u"3、问题解决情况")  # 应用场景示例标题
    run.font.name = u'黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')
    run.font.color.rgb = RGBColor(0, 0, 0)

    #doc.add_heading('3、问题解决情况', level=1)
    table = doc.add_table(rows=1, cols=3)
    row = table.rows[0]
    row.cells[0].text = '问题来源'  # '第1行第一列'
    row.cells[1].text = '问题描述'  # '第1行第二列'
    row.cells[2].text = '是否解决'  # '第1行第三列'
    # 标题4
    run = doc.add_heading('', level=1).add_run(u"4、累计使用情况")  # 应用场景示例标题
    run.font.name = u'黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')
    run.font.color.rgb = RGBColor(0, 0, 0)

    #doc.add_heading('4、累计使用情况', level=1)
    doc.add_paragraph(para3)
    doc.save(r"{}网站侦察兵周报.docx".format(time_))


# 日志统计主函数
# 起始日期 ， 结束日期
log_to_word(20211213000000, 20211218000000)
